"""Tabella servizi della Domanda di Ammissione: l'a-capo con lo
SCONTO/PREZZO RISERVATO compare SOLO nelle righe effettivamente scontate.
Le righe a prezzo pieno restano su una RIGA SINGOLA: prima il <w:br/> del
template era fuori dall'{% if %} e raddoppiava l'altezza di TUTTE le righe.
"""
import pytest
from datetime import date
from decimal import Decimal
from pathlib import Path

from django.conf import settings
from docx import Document as DocxDocument
from docx.oxml.ns import qn


def _righe_dati_tabella_servizi(docx_path):
    """Ritorna le righe dati (no header/totali) della tabella servizi."""
    d = DocxDocument(str(docx_path))
    for t in d.tables:
        header = ' '.join(c.text.strip().lower() for c in t.rows[0].cells)
        if 'descrizione dei servizi' not in header and 'description of the services' not in header:
            continue
        righe = []
        for row in t.rows[1:]:
            qta = row.cells[0].text.strip()
            row_text = ' '.join(c.text.strip().upper() for c in row.cells)
            if (not qta) and (('TOTALE' in row_text) or ('IVA' in row_text)
                              or ('TOTAL' in row_text) or ('VAT' in row_text)):
                continue  # righe totali
            righe.append(row)
        return righe
    return []


@pytest.mark.django_db
def test_acapo_sconto_solo_nelle_righe_scontate(sponsor):
    from sponsors.models import Contact, ContactRole
    from events.models import Event
    from catalog.models import Service
    from contracts.models import Contract, ContractKind, ContractStatus, ContractLine
    from contracts.services.pdf_generator import generate_admission_request_pdf

    Contact.objects.create(
        sponsor=sponsor, full_name='Firma Tario', email='firma@test.it',
        roles=[ContactRole.OPERATIONAL], is_signer=True,
    )
    event = Event.objects.create(
        name={'it': 'Righe Ev', 'en': 'Righe Ev'}, code='RG',
        start_date=date(2026, 11, 1), end_date=date(2026, 11, 2),
    )
    contract = Contract.objects.create(
        sponsor=sponsor, event=event, contract_kind=ContractKind.MAIN,
        status=ContractStatus.SENT, contract_number='RG-26-001',
    )
    svc_pieno = Service.objects.create(
        event=event, code='RG-PIENO', name={'it': 'Spazio espositivo', 'en': 'Space'},
        base_price=Decimal('1000.00'),
    )
    svc_scontato = Service.objects.create(
        event=event, code='RG-SCONTO', name={'it': 'Expert meeting', 'en': 'Expert meeting'},
        base_price=Decimal('5850.00'),
    )
    ContractLine.objects.create(contract=contract, service=svc_pieno, quantity=1)
    ContractLine.objects.create(
        contract=contract, service=svc_scontato, quantity=1,
        discount_amount=Decimal('2500.00'),
    )

    generate_admission_request_pdf(contract)

    docx_path = (Path(settings.MEDIA_ROOT) / 'documents' / 'contracts'
                 / str(contract.id)
                 / f'domanda_ammissione_{contract.contract_number}_{event.id}.docx')
    assert docx_path.exists(), f'docx della domanda mancante: {docx_path}'

    righe = _righe_dati_tabella_servizi(docx_path)
    assert len(righe) == 2, 'attese 2 righe dati nella tabella servizi'

    trovato_sconto = False
    for row in righe:
        cella = row.cells[1]
        testo = cella.text
        brs = cella._tc.findall(f'.//{qn("w:br")}')
        if 'SCONTO RISERVATO' in testo or 'PREZZO RISERVATO' in testo:
            trovato_sconto = True
            assert brs, 'la riga scontata deve avere il testo su due righe'
        else:
            # riga a prezzo pieno: NIENTE a-capo (riga singola)
            assert not brs, f'riga non scontata con a-capo spurio: {testo!r}'
            assert len(cella.paragraphs) == 1
    assert trovato_sconto, 'nessuna riga con sconto trovata nel documento'
