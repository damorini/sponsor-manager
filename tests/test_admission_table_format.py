"""
La tabella servizi della Domanda di Ammissione viene impaginata in modo
professionale dopo il render: font Arial 10pt uniforme, numeri a destra,
descrizione a sinistra, intestazione centrata in grassetto con sfondo.
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _build_sample(path):
    """Ricostruisce la struttura della tabella del template DOPO il render docxtpl:
    intestazione + 1 riga dati (sinistra, 9pt, font ereditato) + righe totali."""
    doc = Document()
    t = doc.add_table(rows=0, cols=4)

    hdr = t.add_row().cells
    for i, txt in enumerate(['q.tà', 'Descrizione dei servizi richiesti', 'Pr. Unit.', 'Euro']):
        run = hdr[i].paragraphs[0].add_run(txt)
        run.font.bold = True

    d = t.add_row().cells
    for i, v in enumerate(['1', 'Spazio espositivo - SD_1', '€ 1.750,00', '€ 1.750,00']):
        run = d[i].paragraphs[0].add_run(v)
        run.font.size = Pt(9)  # come nel template originale

    for label, val in [('TOTALE IMPONIBILE', '1.750,00'), ('IVA 22%', '385,00'), ('TOTALE', '2.135,00')]:
        r = t.add_row().cells
        r[1].paragraphs[0].add_run(label)
        r[3].paragraphs[0].add_run(val)

    doc.save(path)


def test_admission_table_formatting(tmp_path):
    from contracts.services.pdf_generator import _format_admission_services_table

    p = tmp_path / 'sample.docx'
    _build_sample(str(p))
    _format_admission_services_table(str(p))

    doc = Document(str(p))
    t = doc.tables[0]

    # --- riga dati (indice 1): descrizione a sinistra, numeri a destra ---
    data = t.rows[1]
    assert data.cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    for ci in (0, 2, 3):
        assert data.cells[ci].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT

    # --- font uniforme Arial 10pt su tutta la riga dati ---
    for ci in range(4):
        runs = [r for r in data.cells[ci].paragraphs[0].runs if r.text.strip()]
        assert runs, f"cella {ci} senza testo"
        for r in runs:
            assert r.font.name == 'Arial'
            assert r.font.size == Pt(10)

    # --- righe totali a destra in grassetto ---
    for ri in (2, 3, 4):
        tot = t.rows[ri]
        assert tot.cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
        bold_runs = [r for r in tot.cells[1].paragraphs[0].runs if r.text.strip()]
        assert all(r.font.bold for r in bold_runs)

    # --- intestazione: centrata + sfondo grigio ---
    head = t.rows[0]
    assert head.cells[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    from docx.oxml.ns import qn
    shd = head.cells[1]._tc.get_or_add_tcPr().find(qn('w:shd'))
    assert shd is not None and shd.get(qn('w:fill')) == 'D9D9D9'
