"""
Addendum: usa lo stesso modello della Domanda di ammissione, ma con titolo
«ADDENDUM AL CONTRATTO N° xxxx DEL gg/mm/aaaa» (riferito al contratto padre).
"""
import pytest
from datetime import date

from events.models import Event
from contracts.models import Contract, ContractKind, ContractStatus
from contracts.services.pdf_generator import _addendum_title, _replace_title_in_docx


@pytest.mark.django_db
def test_addendum_title_riferisce_il_padre(sponsor):
    event = Event.objects.create(
        name={'it': 'Ev', 'en': 'Ev'}, code='EV',
        start_date=date(2026, 9, 1), end_date=date(2026, 9, 2),
    )
    main = Contract.objects.create(
        sponsor=sponsor, event=event, contract_kind=ContractKind.MAIN,
        status=ContractStatus.SIGNED, contract_number='AB-26-007',
        signed_date=date(2026, 3, 15),
    )
    add = Contract.objects.create(
        sponsor=sponsor, event=event, parent_contract=main,
        contract_kind=ContractKind.ADDENDUM, status=ContractStatus.DRAFT,
        contract_number='AB-26-007-ADD1',
    )
    assert _addendum_title(add) == 'ADDENDUM AL CONTRATTO N° AB-26-007 DEL 15/03/2026'


def test_replace_title_in_docx(tmp_path):
    from docx import Document
    p = tmp_path / 'doc.docx'
    d = Document()
    d.add_paragraph('DOMANDA DI AMMISSIONE')
    d.add_paragraph('Altro testo')
    d.save(str(p))

    _replace_title_in_docx(p, 'ADDENDUM AL CONTRATTO N° X DEL 01/01/2026')

    d2 = Document(str(p))
    assert d2.paragraphs[0].text == 'ADDENDUM AL CONTRATTO N° X DEL 01/01/2026'
    assert d2.paragraphs[1].text == 'Altro testo'  # il resto resta intatto
