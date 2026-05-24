#!/usr/bin/env python3
"""
v2 — Header congresso su TUTTE le pagine + footer 4 righe.

Rispetto alla versione precedente:
  - L'header del congresso va nell'HEADER DI PAGINA (si ripete su ogni pagina),
    non piu' solo in cima alla prima pagina.
  - Rimuove eventuali immagini gia' presenti nell'header di pagina (vecchio
    logo VALET nel template non-ECM), mantenendo il numero pagina.
  - Footer: logo segreteria + 4 righe a bandiera:
        VALET SRL
        <indirizzo>
        Tel: ...  Email: ...  <sito>
        P.IVA ... · REA ...

Sostituisce la funzione _add_header_footer_to_docx gia' presente.
Backup di pdf_generator.py (.bak_pdfv2). Idempotente.

Lancialo dalla cartella del progetto:
    python applica_loghi_pdf_v2.py
"""
import re
import sys
import shutil
from pathlib import Path

ROOT = Path(__file__).resolve().parent
GEN = ROOT / "contracts" / "services" / "pdf_generator.py"
SUFFIX = ".bak_pdfv2"


def fail(msg):
    print(f"\n[X] ERRORE: {msg}")
    print("    Nessuna modifica applicata.")
    sys.exit(1)


if not GEN.exists():
    fail(f"Non trovo {GEN}.")

src = GEN.read_text(encoding="utf-8")

if "HEADER DI PAGINA (v2)" in src:
    print("[OK] La versione v2 sembra gia' presente (salto).")
    sys.exit(0)

if "_add_header_footer_to_docx" not in src:
    fail("Non trovo la funzione _add_header_footer_to_docx. "
         "Va applicata prima la versione base (applica_loghi_pdf.py).")

# Nuova funzione completa
new_func = '''def _add_header_footer_to_docx(docx_path, contract):
    """
    HEADER DI PAGINA (v2): modifica IN-PLACE il .docx generato.

      - HEADER di pagina (su tutte le pagine): inserisce l'immagine header
        dell'evento (contract.event.email_header_image) a tutta larghezza,
        rimuovendo eventuali immagini preesistenti (vecchio logo) ma tenendo
        il testo (es. numero pagina).
      - FOOTER: logo segreteria + 4 righe a bandiera (sinistra) dai dati di
        core.OrganizerSettings.

    Non solleva: ogni parte e' protetta singolarmente.
    """
    from pathlib import Path as _Path
    from docx import Document
    from docx.shared import Mm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    docx_path = _Path(docx_path)
    d = Document(str(docx_path))
    sec = d.sections[0]
    usable_mm = (sec.page_width - sec.left_margin - sec.right_margin) / 36000.0

    changed = False

    # ---- HEADER DI PAGINA: immagine evento su tutte le pagine ----
    event = getattr(contract, 'event', None)
    header_img = getattr(event, 'email_header_image', None) if event else None
    if header_img:
        try:
            img_path = _Path(header_img.path)
            if img_path.exists():
                header = sec.header
                header.is_linked_to_previous = False

                # 1) rimuovi le immagini gia' presenti nell'header (vecchio logo),
                #    tenendo il testo (numero pagina ecc.)
                for p in header.paragraphs:
                    for run in list(p.runs):
                        if '<a:blip' in run._r.xml or '<w:drawing' in run._r.xml:
                            run._r.getparent().remove(run._r)

                # 2) inserisci l'header congresso nel primo paragrafo
                if header.paragraphs:
                    target = header.paragraphs[0]
                else:
                    target = header.add_paragraph()
                target.alignment = WD_ALIGN_PARAGRAPH.CENTER
                target.add_run().add_picture(str(img_path), width=Mm(usable_mm))
                changed = True
        except Exception:
            pass

    # ---- FOOTER: logo segreteria + 4 righe a bandiera ----
    try:
        from core.models import OrganizerSettings
        org = OrganizerSettings.load()
    except Exception:
        org = None

    if org:
        footer = sec.footer
        footer.is_linked_to_previous = False
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)

        # logo (se presente)
        logo = getattr(org, 'logo', None)
        if logo:
            try:
                logo_path = _Path(logo.path)
                if logo_path.exists():
                    p_logo = footer.add_paragraph()
                    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_logo.paragraph_format.space_after = Pt(2)
                    p_logo.add_run().add_picture(str(logo_path), height=Mm(12))
            except Exception:
                pass

        # costruisco le 4 righe (telefono+email+sito sulla stessa riga)
        righe = []
        if org.name:
            righe.append((org.name, True))
        if org.address:
            addr = " ".join(str(org.address).split())  # indirizzo su una riga
            righe.append((addr, False))
        contatti = []
        if org.phone:
            contatti.append("Tel: " + org.phone)
        if org.email:
            contatti.append("Email: " + org.email)
        if org.website:
            contatti.append(org.website)
        if contatti:
            righe.append(("   ".join(contatti), False))
        fisc = []
        if org.vat_number:
            fisc.append("P.IVA " + org.vat_number)
        if org.rea:
            fisc.append("REA " + org.rea)
        if fisc:
            righe.append((" \u00b7 ".join(fisc), False))

        for testo, grassetto in righe:
            pp = footer.add_paragraph()
            pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pp.paragraph_format.space_after = Pt(0)
            pp.paragraph_format.line_spacing = 1.0
            rr = pp.add_run(testo)
            rr.bold = grassetto
            rr.font.size = Pt(7.5)
            rr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        if righe or logo:
            changed = True

    if changed:
        d.save(str(docx_path))
'''

# Sostituisco l'intera vecchia funzione (dalla sua 'def' fino a fine file,
# dato che e' l'ultima funzione del modulo).
pattern = re.compile(
    r"def _add_header_footer_to_docx\(docx_path, contract\):.*\Z",
    re.DOTALL,
)
if not pattern.search(src):
    fail("Non riesco a isolare la vecchia funzione da sostituire.")

new_src = pattern.sub(new_func, src)
if new_src == src:
    fail("Sostituzione non riuscita.")

shutil.copy2(GEN, str(GEN) + SUFFIX)
GEN.write_text(new_src, encoding="utf-8")
print(f"[OK] pdf_generator.py aggiornato a v2 (backup: pdf_generator.py{SUFFIX})")
print("\n=== FATTO. Nessuna migrazione necessaria. ===")
