"""
Generatore PDF dei contratti (versione 2 con template reali VALET).

Differenze rispetto alla v1:
- Salta i contratti ADDON (acquistati dal portale ecommerce, non hanno PDF)
- Sceglie il template giusto in base al tipo evento (ECM vs non-ECM)
- Auto-compila l'allegato con i ContractLine raggruppati per categoria
- Usa i tuoi template VALET originali (con dati anagrafici fissi)

Workflow:
    from contracts.services.pdf_generator import generate_contract_pdf
    
    document = generate_contract_pdf(contract)
    # → documento .docx + .pdf creati e salvati in MEDIA_ROOT
    # → record Document creato con storage_url
    # → ritorna il Document creato (None se è un ADDON)

Templates:
    /home/claude/projects/sponsor_manager/contracts/templates_pdf/
        ├── template_ecm_it.docx
        └── template_non_ecm_it.docx
"""
import logging
from datetime import date
from decimal import Decimal
from collections import defaultdict
from pathlib import Path

from django.conf import settings
from django.contrib.contenttypes.models import ContentType
from django.utils.text import slugify
from django.core.files.storage import default_storage
from django.utils import timezone

from docxtpl import DocxTemplate
from jinja2 import Environment

logger = logging.getLogger(__name__)


# ============================================================================
# Configurazione
# ============================================================================

# Path base dei template (relativo a BASE_DIR di Django)
TEMPLATES_DIR = Path(settings.BASE_DIR) / 'contracts' / 'templates_pdf'

# Mapping tipo evento → template
TEMPLATE_MAP = {
    ('ecm', 'it'): 'template_ecm_it.docx',
    ('non_ecm', 'it'): 'template_non_ecm_it.docx',
    # In futuro: ('ecm', 'en'): 'template_ecm_en.docx', ecc.
}

# Mapping accounting_category → label per allegato ECM
# (ordine di apparizione nell'allegato)
ECM_CATEGORY_LABELS = [
    ('viaggio_partecipanti', 'Spese complessive di viaggio ed ospitalità partecipanti'),
    ('viaggio_relatori', 'Spese di viaggi ed ospitalità relatori'),
    ('affitto_sala', 'Affitto sala'),
    ('stand', 'Spazi espositivi/stand'),
    ('coffee_break', 'Coffee break e colazioni di lavoro'),
    ('scheda_tecnica', 'Scheda tecnica in cartella'),
    ('quota_iscrizione', 'Quota d\u2019iscrizione'),
    ('altro', 'Altre spese'),
]


# ============================================================================
# Filtri Jinja
# ============================================================================

def format_date_filter(value):
    """Formatta una data come dd/mm/yyyy."""
    if not value:
        return ""
    if hasattr(value, 'strftime'):
        return value.strftime("%d/%m/%Y")
    return str(value)


def format_currency_filter(value):
    """Formatta un numero come valuta italiana (1.234,56)."""
    if value is None:
        return "0,00"
    try:
        return f"{float(value):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except (TypeError, ValueError):
        return str(value)


def format_percent_filter(value):
    """Formatta una percentuale senza decimali inutili: 30.00 -> '30', 33.33 -> '33,33'."""
    if value is None:
        return ""
    try:
        f = float(value)
    except (TypeError, ValueError):
        return str(value)
    if f == int(f):
        return str(int(f))
    # decimali utili, virgola italiana, senza zeri finali
    s = f"{f:.2f}".rstrip("0").rstrip(".")
    return s.replace(".", ",")


def get_jinja_env():
    """Restituisce un Jinja Environment con i filtri custom."""
    env = Environment()
    env.filters['format_date'] = format_date_filter
    env.filters['format_currency'] = format_currency_filter
    return env


# ============================================================================
# Generazione PDF principale
# ============================================================================

class _EventForTemplate:
    """Proxy: espone i campi dell'evento ma con .name come stringa localizzata.
    Corregge il caso in cui event.name e' un campo multilingua (dizionario)."""
    def __init__(self, ev):
        self._ev = ev

    @property
    def name(self):
        try:
            return self._ev.get_name() or ''
        except Exception:
            return getattr(self._ev, 'name', '')

    def __getattr__(self, attr):
        return getattr(self._ev, attr)


def _event_for_template(event):
    return _EventForTemplate(event) if event is not None else event


def generate_contract_pdf(contract):
    """
    Genera il PDF (e il .docx intermedio) del contratto.
    
    Args:
        contract: istanza Contract Django
    
    Returns:
        Document creato (instance), oppure None se contract.kind == ADDON
        (per gli addon ecommerce non si genera PDF)
    
    Raises:
        FileNotFoundError se il template non esiste
        ValueError se mancano dati obbligatori (sponsor, evento, firmatario)
    """
    from contracts.models import ContractKind
    from shared.models import Document
    from sponsors.models import Contact

    # ========================================================================
    # 1. SKIP per addon ecommerce
    # ========================================================================
    if contract.contract_kind == ContractKind.ADDON:
        logger.info(
            "Contract %s è un ADDON, skip generazione PDF (ecommerce)",
            contract.contract_number
        )
        return None

    # ========================================================================
    # 2. Scegli template in base a tipo evento e lingua
    # ========================================================================
    event_type = _get_event_type(contract.event)  # 'ecm' o 'non_ecm'
    language = contract.language or 'it'
    
    template_key = (event_type, language)
    template_filename = TEMPLATE_MAP.get(template_key)
    
    if not template_filename:
        # Fallback: non_ecm italiano
        logger.warning(
            "Nessun template per %s, uso fallback non_ecm/it", template_key
        )
        template_filename = TEMPLATE_MAP[('non_ecm', 'it')]
    
    template_path = TEMPLATES_DIR / template_filename
    if not template_path.exists():
        raise FileNotFoundError(
            f"Template contratto non trovato: {template_path}"
        )

    # ========================================================================
    # 3. Costruisci il context per Jinja
    # ========================================================================
    sponsor = contract.sponsor
    event = contract.event
    signer = _get_signer_contact(contract)
    referente = _get_referente_contact(contract)
    
    # Validazione dati minimi
    if not signer:
        raise ValueError(
            f"Contract {contract.contract_number}: nessun Contact con is_signer=True. "
            "Imposta un firmatario sullo Sponsor prima di generare il contratto."
        )
    
    # Raggruppa righe per categoria (per allegato ECM)
    lines_by_category = _group_lines_by_category(contract, event_type)
    
    # Lista servizi piatta (per allegato non-ECM)
    services_list = _build_services_list_string(contract)
    
    context = {
        'contract': contract,
        'sponsor': sponsor,
        'signer': signer,
        'contact_referente': referente,
        'event': _event_for_template(event),
        'lines_by_category': lines_by_category,
        'lines': [ln for _cat, items in lines_by_category for ln in items],  # lista piatta righe (tabella Allegato 2 ECM)
        'services_list': services_list,
        'stand_size': _get_stand_size(contract),
        'signature_place': getattr(settings, 'SIGNATURE_PLACE', 'Bologna'),
        # IVA scorporata (variabili pronte per il template Word del contratto)
        'imponibile': format_currency_filter(contract.subtotal),
        'iva': format_currency_filter(contract.vat_amount),
        'totale': format_currency_filter(contract.total),
        'aliquota_iva': (
            (f"{int(contract.vat_rate)}%" if contract.vat_rate == int(contract.vat_rate)
             else f"{contract.vat_rate}%") if contract.vat_rate else ''
        ),
        'cancellation_penalty_percent': getattr(event, 'cancellation_penalty_percent', 50),
        # Piano pagamento (acconto/saldo), IVA inclusa
        'has_deposit': contract.has_deposit,
        'deposit_percent': contract.deposit_percent,
        'deposit_amount': format_currency_filter(contract.deposit_amount),
        'balance_amount': format_currency_filter(contract.balance_amount),
        'deposit_due_date': format_date_filter(contract.deposit_due_date) if contract.deposit_due_date else '',
        'balance_due_date': format_date_filter(contract.balance_due_date) if contract.balance_due_date else '',
    }
    
    # ========================================================================
    # 4. Render del template .docx
    # ========================================================================
    doc = DocxTemplate(str(template_path))
    doc.render(context, jinja_env=get_jinja_env())
    
    # ========================================================================
    # 5. Salva il .docx intermedio
    # ========================================================================
    docx_filename = f"contratto_{contract.contract_number}_{contract.event.id}.docx"
    relative_docx_path = f"documents/contracts/{contract.id}/{docx_filename}"
    full_docx_path = Path(settings.MEDIA_ROOT) / relative_docx_path
    full_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(full_docx_path))

    # Aggiunge header congresso + footer organizzatore al .docx generato
    # (i template originali restano intatti). Robusto: se manca qualcosa, salta.
    try:
        _add_header_footer_to_docx(full_docx_path, contract)
    except Exception as e:
        logger.warning("Header/footer PDF non applicati per %s: %s",
                       contract.contract_number, e)
    logger.info("Generato .docx: %s", full_docx_path)
    
    # ========================================================================
    # 6. Converti in PDF tramite LibreOffice headless
    # ========================================================================
    pdf_path = _convert_docx_to_pdf(full_docx_path)
    if not pdf_path:
        logger.warning(
            "Conversione PDF fallita per %s, mantengo solo .docx",
            contract.contract_number
        )
        # Crea Document per il .docx
        return _create_document_record(
            contract, full_docx_path, relative_docx_path,
            file_name=docx_filename, mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    # ========================================================================
    # 7. Crea record Document per il PDF
    # ========================================================================
    pdf_filename = pdf_path.name
    relative_pdf_path = relative_docx_path.replace('.docx', '.pdf')
    
    document = _create_document_record(
        contract, pdf_path, relative_pdf_path,
        file_name=pdf_filename, mime='application/pdf'
    )
    
    logger.info(
        "Contract %s: PDF generato e salvato (Document id=%s)",
        contract.contract_number, document.id
    )
    return document


# ============================================================================
# Helper: tipo evento (ECM vs non-ECM)
# ============================================================================

def _get_event_type(event):
    """
    Determina se un evento è ECM o non-ECM.
    
    Cerca un attributo `event.is_ecm` (boolean), oppure `event.event_type`
    ('ecm' / 'non_ecm'), oppure determina da `event.ecm_id`.
    
    Default: non_ecm.
    """
    if hasattr(event, 'is_ecm'):
        return 'ecm' if event.is_ecm else 'non_ecm'
    if hasattr(event, 'event_type') and event.event_type:
        return event.event_type
    if hasattr(event, 'ecm_id') and event.ecm_id:
        return 'ecm'
    return 'non_ecm'


# ============================================================================
# Helper: trova Contact firmatario / referente
# ============================================================================

def _get_signer_contact(contract):
    """
    Trova il firmatario da usare per il contratto, in ordine di priorita':
      1) il "Firmatario" scelto sul contratto (contract.sponsor_signer_contact),
         se valorizzato e non cancellato;
      2) FALLBACK: il primo Contact dello sponsor con is_signer=True
         (comportamento storico, mantenuto invariato).
    """
    chosen = getattr(contract, 'sponsor_signer_contact', None)
    if chosen is not None and getattr(chosen, 'deleted_at', None) is None:
        return chosen
    return contract.sponsor.contacts.filter(
        is_signer=True,
        deleted_at__isnull=True,
    ).first()


def _get_referente_contact(contract):
    """
    Trova il Contact 'referente' dello sponsor (responsabile operativo).
    
    Cerca per role='operational' o is_primary=True.
    """
    contacts = contract.sponsor.contacts.filter(deleted_at__isnull=True)
    
    # Prima prova: contact con ruolo 'operational'
    for c in contacts:
        if 'operational' in (c.roles or []):
            return c
    
    # Fallback: primary contact
    return contacts.filter(is_primary=True).first() or contacts.first()


# ============================================================================
# Helper: raggruppa righe per categoria contabile (per allegato ECM)
# ============================================================================

def _group_lines_by_category(contract, event_type):
    """
    Raggruppa le ContractLine per accounting_category del Service.
    
    Restituisce una LISTA di tuple (category_label, [lines]) ordinata
    secondo ECM_CATEGORY_LABELS. Le categorie senza righe vengono saltate.
    
    Per non-ECM, raggruppa tutto in 'altro' (l'allegato non-ECM ha tabella
    diversa, non usa raggruppamento).
    """
    # Raggruppa per category
    grouped = defaultdict(list)
    for line in contract.lines.select_related('service').all():
        category = (
            getattr(line.service, 'accounting_category', None) or 'altro'
        )
        grouped[category].append(line)
    
    # Ordina secondo ECM_CATEGORY_LABELS (e label friendly)
    result = []
    for cat_key, label in ECM_CATEGORY_LABELS:
        if cat_key in grouped and grouped[cat_key]:
            result.append((label, grouped[cat_key]))
    
    # Aggiungi categorie non standard alla fine
    standard_keys = {k for k, _ in ECM_CATEGORY_LABELS}
    for cat_key, lines in grouped.items():
        if cat_key not in standard_keys and lines:
            result.append((cat_key.replace('_', ' ').capitalize(), lines))
    
    return result


def _build_services_list_string(contract):
    """
    Costruisce una stringa con la lista servizi separati da '; '.
    Usata per la cella 'MODALITà DI SPONSORIZZAZIONE SCELTA' del non-ECM.
    """
    parts = []
    for line in contract.lines.all():
        qty = line.quantity if line.quantity > 1 else None
        if qty:
            parts.append(f"{line.service_name_snapshot} (q.tà {qty})")
        else:
            parts.append(line.service_name_snapshot)
    return "; ".join(parts) if parts else ""


def _get_stand_size(contract):
    """Restituisce la dimensione dello stand assegnato (es. '4x3')."""
    if contract.stand:
        return getattr(contract.stand, 'size_label', None) or contract.stand.code
    if contract.stand_block:
        return contract.stand_block.code
    return None


# ============================================================================
# Helper: conversione docx → pdf via LibreOffice
# ============================================================================

def _convert_docx_to_pdf(docx_path):
    """
    Converte un .docx in PDF usando LibreOffice headless.
    
    Restituisce Path del PDF generato, o None in caso di errore.
    """
    import subprocess
    
    docx_path = Path(docx_path)
    output_dir = docx_path.parent
    
    try:
        result = subprocess.run(
            [
                'libreoffice', '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(output_dir),
                str(docx_path),
            ],
            capture_output=True, timeout=60,
        )
        if result.returncode != 0:
            logger.error(
                "LibreOffice convert error: %s", result.stderr.decode('utf-8', 'ignore')
            )
            return None
    except (subprocess.TimeoutExpired, FileNotFoundError) as e:
        logger.error("LibreOffice non disponibile o timeout: %s", e)
        return None
    
    pdf_path = output_dir / docx_path.name.replace('.docx', '.pdf')
    if pdf_path.exists():
        return pdf_path
    return None


# ============================================================================
# Helper: crea record Document
# ============================================================================

def _create_document_record(contract, full_path, relative_path, file_name, mime, document_type='contract_pdf'):
    """Crea il record Document collegato al contract.

    Anti-doppioni: Soft-elimina i Document precedenti VIVI dello stesso tipo e
    stesso oggetto, cosi' resta sempre un solo documento vivo per tipo (sempre
    aggiornato). I vecchi restano soft-deleted (recuperabili); i file su disco
    non vengono toccati.
    """
    from shared.models import Document
    
    contract_ct = ContentType.objects.get_for_model(contract.__class__)
    
    # Soft-elimina i Document precedenti dello stesso tipo per questo oggetto
    previous = Document.objects.filter(
        content_type=contract_ct,
        object_id=contract.id,
        document_type=document_type,
        deleted_at__isnull=True,
    )
    for old_doc in previous:
        old_doc.delete()  # soft-delete
    
    document = Document.objects.create(
        content_type=contract_ct,
        object_id=contract.id,
        title=file_name,
        file_name=file_name,
        file_size_bytes=full_path.stat().st_size,
        mime_type=mime,
        storage_url=settings.MEDIA_URL + relative_path,
        document_type=document_type,
    )
    return document


# ============================================================================
# Hook: chiamato da Contract.mark_as_sent()
# ============================================================================

def auto_generate_on_send(contract):
    """
    Genera automaticamente il PDF quando il contratto passa a SENT.
    
    Da chiamare alla fine di Contract.mark_as_sent().
    Cattura tutti gli errori per non bloccare la transizione di stato.
    """
    try:
        document = generate_contract_pdf(contract)
        if document:
            logger.info(
                "Auto-generato PDF per %s (Document id=%s)",
                contract.contract_number, document.id
            )
    except Exception as e:
        logger.exception(
            "Errore auto-generazione PDF per %s: %s",
            contract.contract_number, e
        )
        # Non solleviamo: la transizione di stato deve completarsi


# ============================================================================
# Helper: header congresso (immagine) + footer organizzatore nel .docx
# ============================================================================

def _add_header_footer_to_docx(docx_path, contract):
    """
    HEADER DI PAGINA (v2): modifica IN-PLACE il .docx generato.

      - HEADER di pagina (su tutte le pagine): inserisce l'immagine header
        dell'evento (contract.event.email_header_image) a tutta larghezza,
        rimuovendo eventuali immagini preesistenti (vecchio logo) ma tenendo
        il testo (es. numero pagina).
      - FOOTER: logo segreteria + 4 righe a bandiera (sinistra) dai dati di
        core.OrganizerSettings.

    Non solleva: ogni parte e' protetta singolarmente.
    """
    from pathlib import Path as _Path
    from docx import Document
    from docx.shared import Mm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    docx_path = _Path(docx_path)
    d = Document(str(docx_path))
    sec = d.sections[0]
    usable_mm = (sec.page_width - sec.left_margin - sec.right_margin) / 36000.0

    changed = False

    # ---- HEADER DI PAGINA: immagine evento su tutte le pagine ----
    event = getattr(contract, 'event', None)
    header_img = getattr(event, 'email_header_image', None) if event else None
    if header_img:
        try:
            img_path = _Path(header_img.path)
            if img_path.exists():
                header = sec.header
                header.is_linked_to_previous = False

                # 1) rimuovi le immagini gia' presenti nell'header (vecchio logo),
                #    tenendo il testo (numero pagina ecc.)
                for p in header.paragraphs:
                    for run in list(p.runs):
                        if '<a:blip' in run._r.xml or '<w:drawing' in run._r.xml:
                            run._r.getparent().remove(run._r)

                # 2) inserisci l'header congresso nel primo paragrafo
                if header.paragraphs:
                    target = header.paragraphs[0]
                else:
                    target = header.add_paragraph()
                target.alignment = WD_ALIGN_PARAGRAPH.CENTER
                target.add_run().add_picture(str(img_path), width=Mm(usable_mm))
                changed = True
        except Exception:
            pass

    # ---- FOOTER: logo segreteria + 4 righe a bandiera ----
    try:
        from core.models import OrganizerSettings
        org = OrganizerSettings.load()
    except Exception:
        org = None

    if org:
        footer = sec.footer
        footer.is_linked_to_previous = False
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)

        # logo (rilevo il percorso; lo inserisco nella cella di sinistra)
        logo = getattr(org, 'logo', None)
        logo_path = None
        if logo:
            try:
                _lp = _Path(logo.path)
                if _lp.exists():
                    logo_path = _lp
            except Exception:
                logo_path = None
        if logo_path is None:
            try:
                from django.contrib.staticfiles import finders as _finders
                _static_logo = _finders.find('branding/valet_logo.png')
                if _static_logo:
                    logo_path = _Path(_static_logo)
            except Exception:
                pass

        # costruisco le righe di testo (telefono+email+sito sulla stessa riga)
        righe = []
        if org.name:
            righe.append((org.name, True))
        if org.address:
            addr = " ".join(str(org.address).split())  # indirizzo su una riga
            righe.append((addr, False))
        contatti = []
        if org.phone:
            contatti.append("Tel: " + org.phone)
        if org.website:
            contatti.append(org.website)
        if contatti:
            righe.append(("   ".join(contatti), False))
        # Email su una riga propria (etichetta + indirizzo insieme, mai spezzati)
        if org.email:
            righe.append(("Email: " + org.email, False))
        fisc = []
        if org.vat_number:
            fisc.append("P.IVA " + org.vat_number)
        if org.rea:
            fisc.append("REA " + org.rea)
        if fisc:
            righe.append((" \u00b7 ".join(fisc), False))

        # layout affiancato: logo a sinistra, dati a destra (tabella SENZA bordi)
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn as _qn
        logo_w = min(70.0, usable_mm * 0.45)
        text_w = max(60.0, usable_mm - logo_w)
        ftbl = footer.add_table(rows=1, cols=2, width=Mm(usable_mm))
        ftbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        ftbl.autofit = False
        cell_logo, cell_txt = ftbl.rows[0].cells
        cell_logo.width = Mm(logo_w)
        cell_txt.width = Mm(text_w)
        # rimuovo i bordi della tabella
        _tblPr = ftbl._tbl.tblPr
        _bd = _tblPr.makeelement(_qn('w:tblBorders'), {})
        for _edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            _bd.append(_bd.makeelement(_qn('w:' + _edge), {_qn('w:val'): 'none'}))
        _tblPr.append(_bd)
        # cella logo (sinistra)
        p_logo = cell_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if logo_path:
            try:
                p_logo.add_run().add_picture(str(logo_path), height=Mm(12))
            except Exception:
                pass
        # cella dati (destra, a fianco del logo)
        _first = True
        for testo, grassetto in righe:
            pp = cell_txt.paragraphs[0] if _first else cell_txt.add_paragraph()
            _first = False
            pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pp.paragraph_format.space_after = Pt(0)
            pp.paragraph_format.line_spacing = 1.0
            rr = pp.add_run(testo)
            rr.bold = grassetto
            rr.font.name = 'Arial'
            rr.font.size = Pt(7.5)
            rr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        if righe or logo:
            changed = True

    if changed:
        d.save(str(docx_path))


# ============================================================================
# PREVENTIVO: costruzione context segnaposti per il template lettera
# ============================================================================

def build_quote_context(contract):
    """
    Costruisce il dizionario dei segnaposti per il body_template di un
    LetterTemplate, a partire dai dati del contratto.

    Segnaposti prodotti (coerenti con shared.LetterTemplate):
        azienda, evento, date_evento, luogo_evento,
        numero, totale, stand, servizi

    Tutti i valori sono stringhe gia' formattate (date dd/mm/yyyy,
    totale come valuta italiana). Pronto per il render Jinja2.
    """
    sponsor = contract.sponsor
    event = contract.event

    # Nome evento (JSONField multilingua -> stringa nella lingua del contratto)
    lang = getattr(contract, 'language', None) or 'it'
    if hasattr(event, 'get_name'):
        evento = event.get_name(lang)
    else:
        evento = str(event)

    # Date evento (es. "10/06/2026 - 12/06/2026" oppure singola se coincidono)
    start = getattr(event, 'start_date', None)
    end = getattr(event, 'end_date', None)
    if start and end and start != end:
        date_evento = f"{format_date_filter(start)} - {format_date_filter(end)}"
    elif start:
        date_evento = format_date_filter(start)
    else:
        date_evento = ""

    # Luogo evento
    luogo_evento = getattr(event, 'location', "") or ""

    # Stand assegnato (riusa l'helper esistente)
    stand = _get_stand_size(contract) or ""

    # Descrizione stand per il preventivo: override del contratto, altrimenti
    # quella dello stand/blocco assegnato. (Non usata nel contratto.)
    descrizione_stand = (getattr(contract, 'stand_description_override', '') or '').strip()
    if not descrizione_stand:
        _obj = contract.stand or contract.stand_block
        if _obj is not None and hasattr(_obj, 'translated'):
            descrizione_stand = (_obj.translated('quote_description', getattr(contract, 'language', None)) or '').strip()
        else:
            descrizione_stand = ''

    # Frase opzione spazio (solo se il contratto ha un'opzione impostata)
    _opt = getattr(contract, 'option_until', None)
    if _opt:
        opzione_testo = f"Tale spazio risulta opzionato per Voi fino al {_opt.strftime('%d/%m/%Y')}."
    else:
        opzione_testo = ""

    # Lista servizi (riusa l'helper esistente, separati da '; ')
    servizi = _build_services_list_string(contract)

    # Aliquota IVA come stringa leggibile (es. '22%')
    try:
        _rate = contract.vat_rate
        aliquota_iva = f"{int(_rate)}%" if _rate == int(_rate) else f"{_rate}%"
    except Exception:
        aliquota_iva = ""

    context = {
        'azienda': sponsor.legal_name if sponsor else "",
        'evento': evento,
        'date_evento': date_evento,
        'luogo_evento': luogo_evento,
        'numero': contract.contract_number or "",
        'totale': format_currency_filter(contract.total),
        'imponibile': format_currency_filter(contract.subtotal),
        'iva': format_currency_filter(contract.vat_amount),
        'aliquota_iva': aliquota_iva,
        'stand': stand,
        'descrizione_stand': descrizione_stand,
        'opzione_testo': opzione_testo,
        'servizi': servizi,
    }
    return context


def render_quote_body(contract):
    """
    Renderizza il body_template del LetterTemplate associato al contratto,
    compilando i segnaposti con build_quote_context().

    Ritorna la stringa di testo della lettera, oppure '' se non c'e' template.
    """
    template = getattr(contract, 'letter_template', None)
    if not template or not template.body_template:
        return ""
    env = get_jinja_env()
    jinja_template = env.from_string(template.body_template)
    return jinja_template.render(**build_quote_context(contract))


# ============================================================================
# PREVENTIVO: generazione lettera PDF AL VOLO (mattone 5)
# ============================================================================

def generate_quote_pdf(contract):
    """
    Genera la lettera di preventivo in PDF, costruendo il .docx al volo con
    python-docx (nessun template .docx dedicato).

    Struttura della lettera:
        - header di pagina = immagine evento (via _add_header_footer_to_docx)
        - data e luogo (allineati a destra)
        - corpo = testo del LetterTemplate con segnaposti compilati
                  (render_quote_body), un paragrafo per riga
        - footer = dati organizzatore (via _add_header_footer_to_docx)

    Args:
        contract: istanza Contract con letter_template valorizzato.

    Returns:
        Document creato (instance) per il PDF (o per il .docx se la
        conversione PDF fallisce).

    Raises:
        ValueError se manca il LetterTemplate o il corpo risulta vuoto.
    """
    from datetime import date as _date
    from docx import Document as _DocxDocument
    from docx.shared import Pt, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # ----- 1. Validazione: serve un template con corpo -----
    template = getattr(contract, 'letter_template', None)
    if not template:
        raise ValueError(
            f"Contract {contract.contract_number}: nessun template lettera "
            "selezionato. Scegli un 'Template lettera preventivo' sul contratto."
        )

    body = render_quote_body(contract)
    if not body.strip():
        raise ValueError(
            f"Contract {contract.contract_number}: il template '{template.name}' "
            "ha prodotto un corpo vuoto. Controlla il body_template."
        )

    # ----- 2. Costruzione del .docx al volo -----
    doc = _DocxDocument()

    # Margini ragionevoli (lascia spazio a header immagine e footer)
    for section in doc.sections:
        section.top_margin = Mm(38)
        section.bottom_margin = Mm(32)
        section.left_margin = Mm(22)
        section.right_margin = Mm(22)

    # Riga data e luogo (in alto a destra)
    luogo_firma = getattr(settings, 'SIGNATURE_PLACE', 'Bologna')
    event_location = getattr(contract.event, 'location', '') or luogo_firma
    luogo_data = f"{event_location.split(',')[0].strip()}, {format_date_filter(_date.today())}"
    p_data = doc.add_paragraph(luogo_data)
    p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p_data.runs:
        r.font.size = Pt(10)

    doc.add_paragraph("")  # spazio

    # Corpo: un paragrafo per ogni riga del testo renderizzato
    for line in body.split("\n"):
        para = doc.add_paragraph(line)
        para.paragraph_format.space_after = Pt(6)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in para.runs:
            r.font.size = Pt(11)

    # ----- 2b. ALLEGATO 1: riepilogo della proposta -----
    from .client_summary import build_quote_summary_rows
    rows = build_quote_summary_rows(contract)
    if rows:
        doc.add_page_break()
        tit = doc.add_paragraph()
        rt = tit.add_run("Allegato 1 - Riepilogo della proposta")
        rt.bold = True
        rt.font.size = Pt(14)
        doc.add_paragraph("")

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Voce'; hdr[1].text = 'Descrizione'
        hdr[2].text = 'Q.tà'; hdr[3].text = 'Imponibile'
        for rw in rows:
            cells = table.add_row().cells
            cells[0].text = rw['name']
            cells[1].text = rw['description'] or ''
            cells[2].text = str(rw['quantity'])
            cells[3].text = f"€ {format_currency_filter(rw.get('line_subtotal', rw['line_total']))}"

        doc.add_paragraph("")
        # Totali (dal contratto)
        doc.add_paragraph(f"Totale imponibile: € {format_currency_filter(contract.subtotal)}")
        doc.add_paragraph(f"Totale IVA: € {format_currency_filter(contract.vat_amount)}")
        tot_p = doc.add_paragraph()
        tot_r = tot_p.add_run(
            f"Totale complessivo (IVA inclusa): € {format_currency_filter(contract.total)}"
        )
        tot_r.bold = True
        tot_r.font.size = Pt(11)

    # ----- 3. Salva il .docx intermedio -----
    docx_filename = f"preventivo_{contract.contract_number}_{contract.event.id}.docx"
    relative_docx_path = f"documents/quotes/{contract.id}/{docx_filename}"
    full_docx_path = Path(settings.MEDIA_ROOT) / relative_docx_path
    full_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(full_docx_path))

    # ----- 4. Header evento + footer organizzatore (riuso helper esistente) -----
    try:
        _add_header_footer_to_docx(full_docx_path, contract)
    except Exception as e:
        logger.warning("Header/footer preventivo non applicati per %s: %s",
                       contract.contract_number, e)

    logger.info("Generato .docx preventivo: %s", full_docx_path)

    # ----- 5. Conversione in PDF -----
    pdf_path = _convert_docx_to_pdf(full_docx_path)
    if not pdf_path:
        logger.warning("Conversione PDF preventivo fallita per %s, tengo .docx",
                       contract.contract_number)
        return _create_document_record(
            contract, full_docx_path, relative_docx_path,
            file_name=docx_filename,
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            document_type='quote',
        )

    # ----- 6. Record Document per il PDF -----
    pdf_filename = pdf_path.name
    relative_pdf_path = relative_docx_path.replace('.docx', '.pdf')
    document = _create_document_record(
        contract, pdf_path, relative_pdf_path,
        file_name=pdf_filename, mime='application/pdf',
        document_type='quote',
    )
    logger.info("Contract %s: PDF preventivo generato (Document id=%s)",
                contract.contract_number, document.id)
    return document


# ============================================================================
# SCHEDA CLIENTE: PDF riepilogo sponsor+evento (mattone 2)
# ============================================================================

def generate_client_summary_pdf(sponsor, event):
    """
    Genera la scheda cliente in PDF per uno sponsor su un evento.
    Aggrega tutti i contratti non annullati (via build_client_summary),
    costruisce un .docx al volo e lo converte in PDF.

    Returns: Document creato. Raises ValueError se non ci sono contratti.
    """
    from datetime import date as _date
    from docx import Document as _DocxDocument
    from docx.shared import Pt, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from .client_summary import build_client_summary

    data = build_client_summary(sponsor, event)
    if not data['contracts']:
        raise ValueError(
            f"Nessun contratto (non annullato) per {sponsor.legal_name} "
            f"su questo evento: scheda non generabile."
        )

    # contratto di riferimento (per header/footer e per agganciare il Document)
    ref_contract = data['contracts'][0]['contract']
    lang = getattr(ref_contract, 'language', None) or 'it'
    event_name = event.get_name(lang) if hasattr(event, 'get_name') else str(event)

    doc = _DocxDocument()
    for section in doc.sections:
        section.top_margin = Mm(38)
        section.bottom_margin = Mm(32)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)

    # Titolo
    h = doc.add_paragraph()
    r = h.add_run("Scheda cliente")
    r.bold = True
    r.font.size = Pt(16)

    sub = doc.add_paragraph()
    rs = sub.add_run(f"{sponsor.legal_name} - {event_name}")
    rs.font.size = Pt(12)
    doc.add_paragraph("")

    cur = lambda v: format_currency_filter(v)
    dt = lambda d: d.strftime('%d/%m/%Y') if d else '-'

    # Per ogni contratto
    for cd in data['contracts']:
        p = doc.add_paragraph()
        rp = p.add_run(f"Contratto {cd['numero']}"
                       + (f" - Spazio: {cd['stand']}" if cd['stand'] else ""))
        rp.bold = True
        rp.font.size = Pt(12)

        # tabella servizi (con descrizione)
        if cd['lines']:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Servizio'; hdr[1].text = 'Descrizione'
            hdr[2].text = 'Q.tà'; hdr[3].text = 'Imponibile'
            for ln in cd['lines']:
                row = table.add_row().cells
                row[0].text = ln['name']
                row[1].text = ln.get('description', '') or ''
                row[2].text = str(ln['quantity'])
                row[3].text = f"€ {cur(ln.get('line_subtotal', ln['line_total']))}"

        # totali (scorporati e chiari)
        doc.add_paragraph(f"Totale imponibile: € {cur(cd['subtotal'])}")
        doc.add_paragraph(f"Totale IVA: € {cur(cd['vat_amount'])}")
        _pc = doc.add_paragraph()
        _rc = _pc.add_run(f"Totale complessivo (IVA inclusa): € {cur(cd['total'])}")
        _rc.bold = True

        # piano pagamento (con stato scadenze)
        def _stato_label(s):
            return f"  [{s}]" if s else ""
        dep_st = _stato_label(cd.get('deposit_stato'))
        bal_st = _stato_label(cd.get('balance_stato'))
        if cd['has_deposit']:
            doc.add_paragraph(
                f"Acconto ({format_percent_filter(cd['deposit_percent'])}%): € {cur(cd['deposit_amount'])} "
                f"entro il {dt(cd['deposit_due_date'])}{dep_st}  -  "
                f"Saldo: € {cur(cd['balance_amount'])} entro il {dt(cd['balance_due_date'])}{bal_st}"
            )
        else:
            doc.add_paragraph(
                f"Pagamento unico: € {cur(cd['balance_amount'])} "
                f"entro il {dt(cd['balance_due_date'])}{bal_st}"
            )
        doc.add_paragraph("")

    # Riepilogo economico complessivo
    tot = data['totals']
    pay = data['payments']
    rsum = doc.add_paragraph()
    rr = rsum.add_run("Riepilogo")
    rr.bold = True
    rr.font.size = Pt(12)
    doc.add_paragraph(f"Totale contratti (IVA incl.): € {cur(tot['total'])}")
    doc.add_paragraph(f"Incassato: € {cur(pay['incassato'])}")
    doc.add_paragraph(f"Residuo da incassare: € {cur(pay['residuo'])}")

    if pay['movimenti']:
        doc.add_paragraph("")
        mp = doc.add_paragraph(); mr = mp.add_run("Pagamenti ricevuti"); mr.bold = True
        for m in pay['movimenti']:
            doc.add_paragraph(
                f"  - {dt(m['date'])}: € {cur(m['amount'])} ({m['method']}) "
                f"- {m['contract_number']}"
            )

    # Salvataggio docx
    # Nome file leggibile: slug sponsor + slug evento (fallback agli id se vuoti)
    _slug_sponsor = slugify(getattr(sponsor, "legal_name", "") or "") or str(sponsor.id)
    _slug_event = getattr(event, "slug", "") or str(event.id)
    docx_filename = f"scheda_{_slug_sponsor}_{_slug_event}.docx"
    relative_docx_path = f"documents/client_summaries/{sponsor.id}/{docx_filename}"
    full_docx_path = Path(settings.MEDIA_ROOT) / relative_docx_path
    full_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(full_docx_path))

    try:
        _add_header_footer_to_docx(full_docx_path, ref_contract)
    except Exception as e:
        logger.warning("Header/footer scheda cliente non applicati: %s", e)

    pdf_path = _convert_docx_to_pdf(full_docx_path)
    if not pdf_path:
        logger.warning("Conversione PDF scheda cliente fallita, tengo .docx")
        return _create_document_record(
            ref_contract, full_docx_path, relative_docx_path,
            file_name=docx_filename,
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            document_type='client_summary',
        )

    relative_pdf_path = relative_docx_path.replace('.docx', '.pdf')
    document = _create_document_record(
        ref_contract, pdf_path, relative_pdf_path,
        file_name=pdf_path.name, mime='application/pdf',
        document_type='client_summary',
    )
    logger.info("Scheda cliente generata per %s (Document id=%s)",
                sponsor.legal_name, document.id)
    return document


# ============================================================================
# DOMANDA DI AMMISSIONE
# ============================================================================

ADMISSION_TEMPLATE = 'template_domanda_ammissione_it.docx'


def _replace_title_in_docx(docx_path, new_title):
    """Sostituisce il titolo 'DOMANDA DI AMMISSIONE' del docx con new_title,
    preservando lo stile del paragrafo. Usato per gli ADDENDUM (stesso modello,
    titolo diverso)."""
    from docx import Document as _Docx
    d = _Docx(str(docx_path))
    for par in d.paragraphs:
        if par.text.strip().upper() == 'DOMANDA DI AMMISSIONE':
            if par.runs:
                par.runs[0].text = new_title
                for extra in par.runs[1:]:
                    extra.text = ''
            else:
                par.add_run(new_title)
            break
    d.save(str(docx_path))


def _addendum_title(contract):
    """Titolo per l'addendum: «ADDENDUM AL CONTRATTO N° xxxx DEL gg/mm/aaaa»,
    riferito al contratto principale (parent)."""
    parent = contract.parent_contract or contract
    numero = parent.contract_number or contract.contract_number
    data = parent.signed_date or (parent.created_at.date() if parent.created_at else None)
    titolo = f"ADDENDUM AL CONTRATTO N° {numero}"
    if data:
        titolo += f" DEL {data.strftime('%d/%m/%Y')}"
    return titolo


def _format_admission_services_table(docx_path):
    """
    Uniforma la tabella servizi della Domanda di Ammissione DOPO il render docxtpl
    (il template resta intatto). Rende l'impaginazione professionale:
      - font Arial 10pt su TUTTE le celle (uguale al resto del documento; le righe
        dati nel template erano 9pt e con font ereditato, quindi "stonavano");
      - numeri incolonnati a DESTRA (q.tà, Pr. Unit., Euro), descrizione a sinistra;
      - intestazione centrata in grassetto con sfondo grigio chiaro.
    Best-effort e idempotente: se la tabella non c'e', non fa nulla.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BODY_FONT = 'Arial'
    BODY_SIZE = Pt(8)

    def _style_cell(cell, align, bold=None):
        for para in cell.paragraphs:
            para.alignment = align
            for run in para.runs:
                run.font.name = BODY_FONT
                run.font.size = BODY_SIZE
                if bold is not None:
                    run.font.bold = bold

    def _shade(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:shd')):
            tcPr.remove(old)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

    doc = Document(str(docx_path))
    for t in doc.tables:
        header = ' '.join(c.text.strip().lower() for c in t.rows[0].cells)
        if 'descrizione dei servizi' not in header and 'pr. unit' not in header:
            continue
        for ri, row in enumerate(t.rows):
            qta = row.cells[0].text.strip()
            row_text = ' '.join(c.text.strip().upper() for c in row.cells)
            is_total = (not qta) and (('TOTALE' in row_text) or ('IVA' in row_text))
            for ci, cell in enumerate(row.cells):
                if ri == 0:
                    _style_cell(cell, WD_ALIGN_PARAGRAPH.CENTER, bold=True)
                    _shade(cell, 'D9D9D9')
                elif is_total:
                    _style_cell(cell, WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
                else:
                    # riga dati: descrizione (col 1) a sinistra, numeri a destra
                    align = WD_ALIGN_PARAGRAPH.LEFT if ci == 1 else WD_ALIGN_PARAGRAPH.RIGHT
                    _style_cell(cell, align)
        break
    doc.save(str(docx_path))


def generate_admission_request_pdf(contract):
    """
    Genera la DOMANDA DI AMMISSIONE (PDF + .docx) per un contratto/preventivo,
    riusando lo stesso contesto e lo stesso header/footer del contratto.

    Returns:
        Document creato (instance).
    Raises:
        FileNotFoundError se manca il template; ValueError se manca il firmatario.
    """
    from contracts.models import ContractKind
    is_addendum = contract.contract_kind == ContractKind.ADDENDUM

    template_path = TEMPLATES_DIR / ADMISSION_TEMPLATE
    if not template_path.exists():
        raise FileNotFoundError(f"Template domanda non trovato: {template_path}")

    sponsor = contract.sponsor
    event = contract.event
    signer = _get_signer_contact(contract)
    if not signer:
        raise ValueError(
            f"Contract {contract.contract_number}: nessun Contact con is_signer=True. "
            "Imposta un firmatario sullo Sponsor prima di generare la domanda."
        )

    event_type = _get_event_type(event)
    lines_by_category = _group_lines_by_category(contract, event_type)
    lines = [ln for _cat, items in lines_by_category for ln in items]

    # Percentuali "pulite" (40 anziche' 40.00) per il testo del contratto.
    def _pct(v):
        try:
            d = Decimal(str(v))
        except Exception:
            return str(v)
        return str(int(d)) if d == d.to_integral_value() else str(d.normalize())

    has_deposit = bool(contract.has_deposit)
    # Penale cancellazione: se c'e' caparra (acconto+saldo) e' pari alla caparra;
    # se il pagamento e' unico/differito e' la % impostata sull'evento (default 50).
    if has_deposit and contract.deposit_percent:
        _penale = contract.deposit_percent
    else:
        _penale = event.cancellation_penalty_percent or 50

    context = {
        'contract': contract,
        'sponsor': sponsor,
        'signer': signer,
        'event': _event_for_template(event),
        'lines': lines,
        'imponibile': format_currency_filter(contract.subtotal),
        'iva': format_currency_filter(contract.vat_amount),
        'totale': format_currency_filter(contract.total),
        'has_deposit': has_deposit,
        'cancellation_penalty_percent': _pct(event.cancellation_penalty_percent or 50),
        'penale_percent': _pct(_penale),
        'deposit_percent': _pct(contract.deposit_percent or 0),
    }

    doc = DocxTemplate(str(template_path))
    doc.render(context, jinja_env=get_jinja_env())

    file_prefix = 'addendum' if is_addendum else 'domanda_ammissione'
    document_type = 'addendum' if is_addendum else 'admission_request'
    docx_filename = f"{file_prefix}_{contract.contract_number}_{event.id}.docx"
    relative_docx_path = f"documents/contracts/{contract.id}/{docx_filename}"
    full_docx_path = Path(settings.MEDIA_ROOT) / relative_docx_path
    full_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(full_docx_path))

    # Impagina la tabella servizi in modo professionale (font uniforme,
    # numeri a destra, intestazione evidenziata) sul docx generato.
    try:
        _format_admission_services_table(full_docx_path)
    except Exception as e:
        logger.warning("Formattazione tabella domanda non applicata per %s: %s",
                       contract.contract_number, e)

    # Addendum: stesso modello, ma il titolo diventa "ADDENDUM AL CONTRATTO N° ...".
    if is_addendum:
        try:
            _replace_title_in_docx(full_docx_path, _addendum_title(contract))
        except Exception as e:
            logger.warning("Titolo addendum non applicato per %s: %s",
                           contract.contract_number, e)

    try:
        _add_header_footer_to_docx(full_docx_path, contract)
    except Exception as e:
        logger.warning("Header/footer domanda non applicati per %s: %s",
                       contract.contract_number, e)

    pdf_path = _convert_docx_to_pdf(full_docx_path)
    if not pdf_path:
        return _create_document_record(
            contract, full_docx_path, relative_docx_path,
            file_name=docx_filename,
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            document_type=document_type,
        )

    pdf_filename = pdf_path.name
    relative_pdf_path = relative_docx_path.replace('.docx', '.pdf')
    document = _create_document_record(
        contract, pdf_path, relative_pdf_path,
        file_name=pdf_filename, mime='application/pdf',
        document_type=document_type,
    )
    logger.info("Contract %s: documento %s generato (Document id=%s)",
                contract.contract_number, document_type, document.id)
    return document


# ============================================================================
# CONTRATTO DI SPONSORIZZAZIONE (non-ECM) con Domanda di ammissione = Allegato 1
# ============================================================================

SPONSOR_CONTRACT_TEMPLATE = 'template_contratto_sponsor_non_ecm_it.docx'


def _get_operational_contact(contract):
    """Contatto OPERATIVO dello sponsor (destinatario del contratto).
    Fallback: referente/primary (via _get_referente_contact) e poi firmatario."""
    ref = _get_referente_contact(contract)
    return ref or _get_signer_contact(contract)


def generate_sponsor_contract_pdf(contract):
    """Genera il CONTRATTO DI SPONSORIZZAZIONE (non-ECM) in PDF, compilato coi
    dati di sponsor/firmatario/evento, con la DOMANDA DI AMMISSIONE allegata
    come ALLEGATO 1 (ultime pagine) in un unico PDF.

    Ritorna il Document creato (document_type='sponsor_contract').
    """
    template_path = TEMPLATES_DIR / SPONSOR_CONTRACT_TEMPLATE
    if not template_path.exists():
        raise FileNotFoundError(f"Template contratto sponsor non trovato: {template_path}")

    sponsor = contract.sponsor
    event = contract.event
    signer = _get_signer_contact(contract)
    if not signer:
        raise ValueError(
            f"Contract {contract.contract_number}: manca il firmatario, "
            "impossibile generare il contratto di sponsorizzazione."
        )
    ref = _get_operational_contact(contract)
    operational_email = (getattr(ref, 'email', '') or getattr(signer, 'email', '') or '')

    context = {
        'contract': contract,
        'sponsor': sponsor,
        'signer': signer,
        'event': _event_for_template(event),
        'operational_email': operational_email,
    }
    doc = DocxTemplate(str(template_path))
    doc.render(context, jinja_env=get_jinja_env())

    docx_filename = f"contratto_sponsor_{contract.contract_number}_{event.id}.docx"
    relative_docx_path = f"documents/contracts/{contract.id}/{docx_filename}"
    full_docx_path = Path(settings.MEDIA_ROOT) / relative_docx_path
    full_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(full_docx_path))

    try:
        _add_header_footer_to_docx(full_docx_path, contract)
    except Exception as e:
        logger.warning("Header/footer contratto sponsor non applicati per %s: %s",
                       contract.contract_number, e)

    contract_pdf = _convert_docx_to_pdf(full_docx_path)
    if not contract_pdf:
        return _create_document_record(
            contract, full_docx_path, relative_docx_path, file_name=docx_filename,
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            document_type='sponsor_contract',
        )

    # ALLEGATO 1: la Domanda di ammissione (riusa quella gia' generata, se c'e')
    domanda_pdf = None
    try:
        from shared.models import Document
        ct = ContentType.objects.get_for_model(contract.__class__)
        adm = (Document.objects
               .filter(content_type=ct, object_id=contract.id,
                       document_type='admission_request', deleted_at__isnull=True)
               .order_by('-created_at').first())
        if not (adm and str(adm.storage_url).endswith('.pdf')):
            adm = generate_admission_request_pdf(contract)
        if adm and str(adm.storage_url).endswith('.pdf'):
            rel = adm.storage_url.replace(settings.MEDIA_URL, '', 1)
            p = Path(settings.MEDIA_ROOT) / rel
            if p.exists():
                domanda_pdf = p
    except Exception as e:
        logger.warning("Domanda (Allegato 1) non disponibile per %s: %s",
                       contract.contract_number, e)

    final_pdf = contract_pdf
    final_name = contract_pdf.name
    if domanda_pdf:
        try:
            from pypdf import PdfWriter
            merged_name = f"contratto_sponsor_completo_{contract.contract_number}_{event.id}.pdf"
            merged_path = full_docx_path.parent / merged_name
            writer = PdfWriter()
            writer.append(str(contract_pdf))
            writer.append(str(domanda_pdf))
            with open(merged_path, 'wb') as fh:
                writer.write(fh)
            writer.close()
            final_pdf = merged_path
            final_name = merged_name
        except Exception as e:
            logger.warning("Merge contratto+domanda (Allegato 1) fallito per %s: %s",
                           contract.contract_number, e)

    relative_pdf_path = f"documents/contracts/{contract.id}/{final_name}"
    document = _create_document_record(
        contract, final_pdf, relative_pdf_path, file_name=final_name,
        mime='application/pdf', document_type='sponsor_contract',
    )
    logger.info("Contract %s: contratto sponsor generato (Document id=%s)",
                contract.contract_number, document.id)
    return document


# ============================================================================
# PREVENTIVO: PDF dalla grafica HTML (stessa della mail) con link cliccabile
# ============================================================================

def build_scientific_secretariat_context(event, site_url=''):
    """Costruisce il dict {'text', 'logo_url'} della Segreteria Scientifica per
    il footer del preventivo, dai campi Event.scientific_secretariat[_logo].

    Ritorna None se non c'e' ne' testo ne' logo (cosi' il template non stampa
    nulla a destra)."""
    site_url = (site_url or '').rstrip('/')
    text = (getattr(event, 'scientific_secretariat', '') or '').strip()
    logo_url = ''
    try:
        logo = getattr(event, 'scientific_secretariat_logo', None)
        if logo:
            logo_url = (site_url + logo.url) if site_url else logo.url
    except Exception:
        logo_url = ''
    if not text and not logo_url:
        return None
    return {'text': text, 'logo_url': logo_url}


def generate_quote_pdf_html(contract):
    """Genera il PDF del preventivo dalla grafica HTML (come la mail), con
    pulsante cliccabile verso la pagina del portale. Richiede WeasyPrint."""
    from django.template.loader import render_to_string
    from django.urls import reverse
    try:
        from weasyprint import HTML as _WeasyHTML
    except Exception as e:
        raise RuntimeError(
            "WeasyPrint non installato: esegui 'pip install weasyprint' "
            "(e le librerie di sistema pango/cairo). Dettaglio: %s" % e
        )

    sponsor = contract.sponsor
    event = contract.event
    lines = list(contract.lines.all())
    site_url = getattr(settings, 'SITE_URL', '').rstrip('/')
    try:
        portal_path = reverse('portal:contract_detail', args=[contract.id])
    except Exception:
        portal_path = f"/contracts/{contract.id}/"
    header_url = ''
    try:
        _img = getattr(event, 'email_header_image', None)
        if _img:
            header_url = (site_url + _img.url) if site_url else _img.url
    except Exception:
        header_url = ''
    # Validita' del preventivo: opzione fino alla data impostata sul contratto,
    # altrimenti 30 giorni dalla data di creazione.
    from datetime import timedelta
    _created = getattr(contract, 'created_at', None)
    _created_date = _created.date() if _created else None
    quote_valid_until = contract.option_until or (
        (_created_date + timedelta(days=30)) if _created_date else None)

    # Dati organizzatore (segreteria) per il footer del preventivo.
    try:
        from core.models import OrganizerSettings
        _org = OrganizerSettings.load()
    except Exception:
        _org = None
    org_logo_url = ''
    try:
        if _org and _org.logo:
            org_logo_url = (site_url + _org.logo.url) if site_url else _org.logo.url
        else:
            # Nessun logo caricato in OrganizerSettings: usa il logo brand VALET
            # dagli static (come fa il footer delle email). Per WeasyPrint
            # preferiamo il path su disco (file://), cosi' compare SEMPRE nel PDF
            # senza dipendere dalla raggiungibilita' di rete; in fallback l'URL
            # statico assoluto.
            from django.contrib.staticfiles import finders as _finders
            _logo_fs = _finders.find('branding/valet_logo.png')
            if _logo_fs:
                org_logo_url = Path(_logo_fs).as_uri()
            else:
                from django.templatetags.static import static as _static
                org_logo_url = (site_url + _static('branding/valet_logo.png')) if site_url else _static('branding/valet_logo.png')
    except Exception:
        org_logo_url = ''
    org = {
        'name': (_org.name if _org and _org.name else None)
                or getattr(settings, 'ORGANIZER_DISPLAY_NAME', 'VALET S.r.l.'),
        'address': (_org.address if _org else '') or '',
        'email': (_org.email if _org else '') or getattr(settings, 'SUPPORT_EMAIL', ''),
        'phone': (_org.phone if _org else '') or '',
        'website': (_org.website if _org else '') or '',
        'vat': (_org.vat_number if _org else '') or '',
        'rea': (_org.rea if _org else '') or '',
        'logo_url': org_logo_url,
    }

    # Segreteria Scientifica (per-evento): footer in basso a destra. None se assente.
    sci = build_scientific_secretariat_context(event, site_url)

    # ---- Testi del preventivo nella lingua del contratto (IT/EN) ----
    from django.utils.html import escape, mark_safe
    lang = (getattr(contract, 'language', '') or 'it')
    EN = lang == 'en'
    try:
        ev_name = event.get_name(lang) if hasattr(event, 'get_name') else str(event)
    except Exception:
        ev_name = str(event)
    sp_name = sponsor.legal_name if sponsor else ''
    luogo = getattr(event, 'venue_name', '') or getattr(event, 'location', '') or ''
    data_ev = ''
    if getattr(event, 'start_date', None):
        data_ev = event.start_date.strftime('%d/%m/%Y')
        _end = getattr(event, 'end_date', None)
        if _end and _end != event.start_date:
            data_ev += ' – ' + _end.strftime('%d/%m/%Y')
    valid_s = quote_valid_until.strftime('%d/%m/%Y') if quote_valid_until else ''
    _sp, _ev, _lu = escape(sp_name), escape(ev_name), escape(luogo)

    if EN:
        intro = mark_safe(
            f"Dear <strong>{_sp}</strong>,<br>following the preferences you indicated, we are "
            f"pleased to submit our offer for your company's presence at <strong>{_ev}</strong>"
            + (f", to be held in {_lu}" if luogo else "")
            + (f" on {data_ev}" if data_ev else "")
            + ". We are confident that these solutions can be an excellent showcase for your "
            "company. The space and services selected are summarised below; please take careful "
            "note of the details.")
        attn = ("This document is a quote. The spaces indicated are not currently booked, but "
                "simply held as an option" + (f" until {valid_s}" if valid_s else "") + ".")
        validity = mark_safe(f"This quote is valid until <strong>{valid_s}</strong>. The spaces "
                             "are not booked yet, but simply held as an option.")
        t = {
            'eyebrow': 'Sponsorship proposal', 'intro': intro, 'attn': attn,
            'section_title': 'Summary of spaces and services on option',
            'empty': 'No items selected.', 'incl': 'Included',
            'sconto': 'Unconditional discount',
            'imponibile': 'Net amount', 'iva': 'VAT', 'totale': 'Total (VAT included)',
            'validity': validity, 'cta': 'View and confirm the quote', 'ref': 'Quote',
        }
    else:
        intro = mark_safe(
            f"Spett.le <strong>{_sp}</strong>,<br>a seguito delle preferenze da voi indicate, siamo "
            f"lieti di sottoporvi l'offerta relativa alla presenza della vostra azienda nell'ambito "
            f"di <strong>{_ev}</strong>"
            + (f", che si terrà a {_lu}" if luogo else "")
            + (f" in data {data_ev}" if data_ev else "")
            + ". Siamo certi che queste soluzioni potranno rappresentare un'ottima vetrina per la "
            "vostra azienda. Lo spazio e i servizi selezionati sono riassunti qui di seguito; vi "
            "preghiamo di prendere buona nota di quanto indicato.")
        attn = ("Il presente documento è un preventivo. Gli spazi indicati non risultano al "
                "momento prenotati, ma semplicemente opzionati"
                + (f" fino al {valid_s}" if valid_s else "") + ".")
        validity = mark_safe(f"Questo preventivo è valido fino al <strong>{valid_s}</strong>. "
                             "Al momento gli spazi non sono prenotati ma semplicemente opzionati.")
        t = {
            'eyebrow': 'Proposta di sponsorizzazione', 'intro': intro, 'attn': attn,
            'section_title': 'Riepilogo spazi e servizi in opzione',
            'empty': 'Nessuna voce selezionata.', 'incl': 'Incluso',
            'sconto': 'Sconto incondizionato',
            'imponibile': 'Imponibile', 'iva': 'IVA', 'totale': 'Totale (IVA inclusa)',
            'validity': validity, 'cta': 'Vedi e conferma il preventivo', 'ref': 'Preventivo',
        }

    from decimal import Decimal as _Dec
    sconto_totale = sum((getattr(l, 'sconto_eur', _Dec('0')) or _Dec('0')) for l in lines)
    ctx = {
        'contract': contract,
        'sponsor': sponsor,
        'event': _event_for_template(event),
        'lines': lines,
        'sconto_totale': sconto_totale,
        'confirm_url': (site_url + portal_path) if site_url else portal_path,
        'header_url': header_url,
        'brand_color': getattr(settings, 'BRAND_PRIMARY_COLOR', '#1d6534'),
        'quote_valid_until': quote_valid_until,
        'org': org,
        'sci': sci,
        't': t,
        'ev_name': ev_name,
    }
    html = render_to_string('quote_pdf.html', ctx)
    pdf_bytes = _WeasyHTML(string=html, base_url=(site_url or None)).write_pdf()

    pdf_filename = f"preventivo_{contract.contract_number}_{event.id}.pdf"
    relative_pdf_path = f"documents/contracts/{contract.id}/{pdf_filename}"
    full_pdf_path = Path(settings.MEDIA_ROOT) / relative_pdf_path
    full_pdf_path.parent.mkdir(parents=True, exist_ok=True)
    full_pdf_path.write_bytes(pdf_bytes)

    document = _create_document_record(
        contract, full_pdf_path, relative_pdf_path,
        file_name=pdf_filename, mime='application/pdf', document_type='quote',
    )
    logger.info("Preventivo (HTML->PDF) generato per %s (Document id=%s)",
                contract.contract_number, document.id)
    return document
