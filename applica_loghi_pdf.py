#!/usr/bin/env python3
"""
Aggiunge al generatore PDF dei contratti:
  - HEADER: immagine del congresso (event.email_header_image) in cima al corpo,
    a tutta larghezza utile della pagina.
  - FOOTER: logo segreteria + dati organizzatore (da core.OrganizerSettings),
    su piu' righe allineate a sinistra (a bandiera).

I template Word ORIGINALI non vengono toccati: la modifica avviene sul .docx
generato al momento, prima della conversione in PDF.
Se manca l'header o i dati segreteria, quella parte viene semplicemente saltata
(nessun errore).

Backup di pdf_generator.py (.bak_pdfloghi). Idempotente.

Lancialo dalla cartella del progetto:
    python applica_loghi_pdf.py
"""
import sys
import shutil
from pathlib import Path

ROOT = Path(__file__).resolve().parent
GEN = ROOT / "contracts" / "services" / "pdf_generator.py"
SUFFIX = ".bak_pdfloghi"


def fail(msg):
    print(f"\n[X] ERRORE: {msg}")
    print("    Nessuna modifica applicata.")
    sys.exit(1)


if not GEN.exists():
    fail(f"Non trovo {GEN}.")

src = GEN.read_text(encoding="utf-8")

if "_add_header_footer_to_docx" in src:
    print("[OK] La funzione header/footer sembra gia' presente (salto).")
    sys.exit(0)

# --- 1. Inserisco la chiamata dopo doc.save(...) del docx intermedio ---
anchor = '    doc.save(str(full_docx_path))\n'
if anchor not in src:
    fail("Non trovo 'doc.save(str(full_docx_path))' in pdf_generator.py.")

call_block = (
    '    doc.save(str(full_docx_path))\n'
    '\n'
    '    # Aggiunge header congresso + footer organizzatore al .docx generato\n'
    '    # (i template originali restano intatti). Robusto: se manca qualcosa, salta.\n'
    '    try:\n'
    '        _add_header_footer_to_docx(full_docx_path, contract)\n'
    '    except Exception as e:\n'
    '        logger.warning("Header/footer PDF non applicati per %s: %s",\n'
    '                       contract.contract_number, e)\n'
)
src = src.replace(anchor, call_block, 1)

# --- 2. Aggiungo la funzione in fondo al file ---
func_code = '''

# ============================================================================
# Helper: header congresso (immagine) + footer organizzatore nel .docx
# ============================================================================

def _add_header_footer_to_docx(docx_path, contract):
    """
    Modifica IN-PLACE il .docx generato:
      - inserisce in cima al corpo l'immagine header dell'evento, se presente
        (contract.event.email_header_image);
      - imposta il footer con logo segreteria + dati organizzatore presi da
        core.OrganizerSettings.

    Non solleva: ogni parte e' protetta singolarmente.
    """
    from pathlib import Path as _Path
    from docx import Document
    from docx.shared import Mm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    docx_path = _Path(docx_path)
    d = Document(str(docx_path))
    sec = d.sections[0]
    usable_mm = (sec.page_width - sec.left_margin - sec.right_margin) / 36000.0

    changed = False

    # ---- HEADER: immagine evento in cima al corpo ----
    event = getattr(contract, 'event', None)
    header_img = getattr(event, 'email_header_image', None) if event else None
    if header_img:
        try:
            img_path = _Path(header_img.path)
            if img_path.exists():
                body = d.element.body
                new_p = body.makeelement(qn('w:p'), {})
                body.insert(0, new_p)
                para = Paragraph(new_p, d)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run().add_picture(str(img_path), width=Mm(usable_mm))
                changed = True
        except Exception:
            pass

    # ---- FOOTER: logo segreteria + dati a bandiera ----
    try:
        from core.models import OrganizerSettings
        org = OrganizerSettings.load()
    except Exception:
        org = None

    if org:
        footer = sec.footer
        footer.is_linked_to_previous = False
        # svuota footer esistente
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)

        # logo (se presente)
        logo = getattr(org, 'logo', None)
        if logo:
            try:
                logo_path = _Path(logo.path)
                if logo_path.exists():
                    p_logo = footer.add_paragraph()
                    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_logo.paragraph_format.space_after = Pt(2)
                    p_logo.add_run().add_picture(str(logo_path), height=Mm(12))
            except Exception:
                pass

        # righe dati (testo, grassetto)
        righe = []
        if org.name:
            righe.append((org.name, True))
        if org.address:
            for ln in str(org.address).splitlines():
                if ln.strip():
                    righe.append((ln.strip(), False))
        if org.phone:
            righe.append(("Tel: " + org.phone, False))
        if org.email:
            righe.append(("Email: " + org.email, False))
        if org.website:
            righe.append((org.website, False))
        fisc = []
        if org.vat_number:
            fisc.append("P.IVA " + org.vat_number)
        if org.rea:
            fisc.append("REA " + org.rea)
        if fisc:
            righe.append((" \u00b7 ".join(fisc), False))

        for testo, grassetto in righe:
            pp = footer.add_paragraph()
            pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pp.paragraph_format.space_after = Pt(0)
            pp.paragraph_format.line_spacing = 1.0
            rr = pp.add_run(testo)
            rr.bold = grassetto
            rr.font.size = Pt(7.5)
            rr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        if righe or logo:
            changed = True

    if changed:
        d.save(str(docx_path))
'''

src = src.rstrip() + "\n" + func_code

shutil.copy2(GEN, str(GEN) + SUFFIX)
GEN.write_text(src, encoding="utf-8")
print(f"[OK] pdf_generator.py aggiornato (backup: pdf_generator.py{SUFFIX})")
print("\n=== FATTO. Nessuna migrazione necessaria (solo codice). ===")
